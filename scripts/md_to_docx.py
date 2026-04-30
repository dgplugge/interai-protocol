"""md_to_docx.py — minimal Markdown → DOCX converter for project status docs.

Handles the subset we actually use in state-of-project readouts:
- ATX headings (# ## ###)
- Paragraphs with **bold**, *italic*, `inline code`, [link](url)
- Bullet lists (- item)
- Numbered lists (1. item)
- Pipe tables with header separator row
- Block quotes (> ...)
- Horizontal rules (---)

Not a full Markdown engine. Good enough for our internal status docs.

Usage:
    python scripts/md_to_docx.py <input.md> <output.docx> [--title "Document Title"]
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
NUMBERED_RE = re.compile(r"^(\s*)\d+\.\s+(.*)$")
TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|[\s|:\-]+\|\s*$")
HR_RE = re.compile(r"^\s*-{3,}\s*$")
BLOCKQUOTE_RE = re.compile(r"^\s*>\s?(.*)$")

INLINE_CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def _add_runs_with_inline_formatting(paragraph, text: str) -> None:
    """Walk inline `code`, **bold**, *italic*, [link](url) and emit runs."""
    pos = 0
    pattern = re.compile(
        r"(`[^`]+`|\*\*[^*]+\*\*|(?<!\*)\*[^*\n]+\*(?!\*)|\[[^\]]+\]\([^)]+\))"
    )
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("**") and token.endswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*") and token.endswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        elif token.startswith("["):
            link_m = LINK_RE.match(token)
            if link_m:
                paragraph.add_run(link_m.group(1)).font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "BFBFBF")
        tblBorders.append(b)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _set_cell_shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _emit_table(doc, header_row: list[str], data_rows: list[list[str]]) -> None:
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    _set_table_borders(table)

    # Header
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        para = cell.paragraphs[0]
        _add_runs_with_inline_formatting(para, cell_text)
        for run in para.runs:
            run.bold = True
        _set_cell_shading(cell, "E7EEF6")

    for ri, row in enumerate(data_rows, start=1):
        for ci in range(cols):
            text = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            para = cell.paragraphs[0]
            _add_runs_with_inline_formatting(para, text)


def md_to_docx(md_path: Path, docx_path: Path, title: str | None) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    # Set base font / margins
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    if title:
        t = doc.add_paragraph()
        run = t.add_run(title)
        run.bold = True
        run.font.size = Pt(20)
        t.paragraph_format.space_after = Pt(12)

    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        if HR_RE.match(line):
            # Horizontal rule — render as a thin paragraph border
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            top = OxmlElement("w:bottom")
            top.set(qn("w:val"), "single")
            top.set(qn("w:sz"), "6")
            top.set(qn("w:space"), "1")
            top.set(qn("w:color"), "BFBFBF")
            pBdr.append(top)
            pPr.append(pBdr)
            i += 1
            continue

        m = HEADING_RE.match(line)
        if m:
            level = min(len(m.group(1)), 4)
            text = m.group(2).strip()
            heading = doc.add_heading(level=level)
            _add_runs_with_inline_formatting(heading, text)
            i += 1
            continue

        # Table block
        if TABLE_ROW_RE.match(line):
            block = []
            while i < len(lines) and TABLE_ROW_RE.match(lines[i]):
                block.append(lines[i])
                i += 1
            if len(block) >= 2 and TABLE_SEP_RE.match(block[1]):
                header = _split_table_row(block[0])
                data = [_split_table_row(r) for r in block[2:]]
                _emit_table(doc, header, data)
            else:
                # No separator — treat each row as a paragraph
                for r in block:
                    p = doc.add_paragraph()
                    _add_runs_with_inline_formatting(p, r)
            continue

        # Bullet list block
        if BULLET_RE.match(line):
            while i < len(lines) and BULLET_RE.match(lines[i]):
                m2 = BULLET_RE.match(lines[i])
                p = doc.add_paragraph(style="List Bullet")
                _add_runs_with_inline_formatting(p, m2.group(2).strip())
                i += 1
            continue

        # Numbered list block
        if NUMBERED_RE.match(line):
            while i < len(lines) and NUMBERED_RE.match(lines[i]):
                m2 = NUMBERED_RE.match(lines[i])
                p = doc.add_paragraph(style="List Number")
                _add_runs_with_inline_formatting(p, m2.group(2).strip())
                i += 1
            continue

        # Block quote
        if BLOCKQUOTE_RE.match(line):
            chunks = []
            while i < len(lines) and BLOCKQUOTE_RE.match(lines[i]):
                m2 = BLOCKQUOTE_RE.match(lines[i])
                chunks.append(m2.group(1))
                i += 1
            joined = " ".join(chunks)
            p = doc.add_paragraph(style="Intense Quote")
            _add_runs_with_inline_formatting(p, joined)
            continue

        # Default: paragraph (may span multiple lines until blank)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not (
            HEADING_RE.match(lines[i])
            or BULLET_RE.match(lines[i])
            or NUMBERED_RE.match(lines[i])
            or TABLE_ROW_RE.match(lines[i])
            or HR_RE.match(lines[i])
            or BLOCKQUOTE_RE.match(lines[i])
        ):
            para_lines.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        _add_runs_with_inline_formatting(p, " ".join(s.strip() for s in para_lines))

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("input", type=Path)
    ap.add_argument("output", type=Path)
    ap.add_argument("--title", default=None)
    args = ap.parse_args()
    md_to_docx(args.input, args.output, args.title)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
